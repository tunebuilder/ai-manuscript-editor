"""
Script to create test fixtures (PDF and DOCX) from sample journal text.
Run this script to generate test files for unit testing.
"""

from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from docx import Document
from docx.shared import Inches
import os


def create_test_pdf(text_content: str, output_path: str):
    """Create a test PDF file from text content."""
    doc = SimpleDocTemplate(output_path, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    
    # Split content into paragraphs
    paragraphs = text_content.split('\n\n')
    
    for para in paragraphs:
        if para.strip():
            # Check if it's a day header
            if para.strip().startswith('Day '):
                # Use heading style for day headers
                p = Paragraph(para.strip(), styles['Heading2'])
            else:
                # Use normal style for content
                p = Paragraph(para.strip(), styles['Normal'])
            
            story.append(p)
            story.append(Spacer(1, 12))  # Add some space
    
    doc.build(story)
    print(f"Created test PDF: {output_path}")


def create_test_docx(text_content: str, output_path: str):
    """Create a test DOCX file from text content."""
    doc = Document()
    
    # Split content into paragraphs
    paragraphs = text_content.split('\n\n')
    
    for para in paragraphs:
        if para.strip():
            # Check if it's a day header
            if para.strip().startswith('Day '):
                # Add as heading
                doc.add_heading(para.strip(), level=2)
            else:
                # Add as normal paragraph
                doc.add_paragraph(para.strip())
    
    doc.save(output_path)
    print(f"Created test DOCX: {output_path}")


def main():
    """Create test fixtures from sample journal text."""
    # Create test_fixtures directory if it doesn't exist
    os.makedirs('test_fixtures', exist_ok=True)
    
    # Read sample journal text
    with open('test_fixtures/sample_journal.txt', 'r', encoding='utf-8') as f:
        text_content = f.read()
    
    # Create test files
    create_test_pdf(text_content, 'test_fixtures/sample_journal.pdf')
    create_test_docx(text_content, 'test_fixtures/sample_journal.docx')
    
    print("Test fixtures created successfully!")


if __name__ == "__main__":
    main() 